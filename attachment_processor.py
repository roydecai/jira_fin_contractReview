import io
import json
import pdfplumber
from docx import Document


def conver_pdf_response_to_json(response):
    # 判断是否为pdf文档
    content_type = response.headers.get("Content-Type","")
    PDF_MIME = "application/pdf"

    if PDF_MIME in content_type:
        print("确认文件类型: PDF，开始解析...")
        try:
            pdf_text = ""
            with pdfplumber.open(io.BytesIO(response.content)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    pdf_text += f"=== 第{page_num}页 ===\n"
                    pdf_text += page_text + "\n\n"

                # 封装成豆包API要求的json结构
                result_json = {
                    "text": pdf_text,
                    "pdf_info": {
                        "total_pages": len(pdf.pages)
                    }
                }

            # 生成json文件，并返回json字符串
            json_str = json.dumps(result_json, ensure_ascii=False, indent=4)
            return json_str

        except Exception as e:
            return  json.dumps({"error": f"PDF处理失败：{str(e)}"}, ensure_ascii=False)
    else:
        return json.dumps({"error": "非PDF格式文件，无法处理"}, ensure_ascii=False)


def conver_docx_response_to_json(response):
    # 0. 判断是否为docx文档
    content_type = response.headers.get("Content-Type","")
    DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    DOC_MIME = "application/msword"

    if DOCX_MIME in content_type:
        print("确认文件类型: Word (.docx)，开始解析...")
    elif DOC_MIME in content_type:
        print("检测到旧版 .doc 格式文件，请转换为 .docx 格式文件上传")
        return json.dumps({"error": "不支持.doc格式，仅支持.docx格式"}, ensure_ascii=False)
    else:
        error_msg = f"附件格式为：{content_type}，程序不支持。请上传 .docx 或 PDF 格式文件！"
        print(error_msg)
        return json.dumps({"error": error_msg}, ensure_ascii=False)


    # 1. 将二进制流放入内存文件对象
    doc_file = io.BytesIO(response.content)

    # 2. 加载文档
    try:
        doc = Document(doc_file)
        full_content = {
            "title": "Jira Attachment Content",
            "paragraphs": [],
            "tables": []
        }

        # 3. 提取段落文字
        for para in doc.paragraphs:
            if para.text.strip():
                full_content["paragraphs"].append(para.text.strip())

        # 4. 提取表格信息
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells]
                table_data.append(row_content)
            full_content["tables"].append(table_data)

        # 5. 转换为JSON格式
        return json.dumps(full_content, ensure_ascii=False, indent=2)

    except Exception as e:
        return json.dumps({"error": f"DOCX处理失败：{str(e)}"}, ensure_ascii=False)


def convert_response_to_json(response):

    if not response or not hasattr(response, "headers"):
        return json.dumps({"error": "无效的响应对象"}, ensure_ascii=False)

    response_type = response.headers.get("Content-Type", "")
    PDF_MIME = "application/pdf"
    DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if PDF_MIME in response_type:
        return conver_pdf_response_to_json(response)
    elif DOCX_MIME in response_type:
        return conver_docx_response_to_json(response)
    else:
        return json.dumps({"error": f"不支持的文件格式：{response_type}"}, ensure_ascii=False)



# 程序测试
if __name__ == "__main__":
    import jira_client

    try:
        # 1. 初始化客户端（官方鉴权）
        jira_cli = jira_client.connect_jira()
        # 2. 获取目标ticket（官方POST/search）
        issues = jira_cli.get_target_issues()
        # 3. 验证结果
        if issues:
            print(f"✅ 测试成功！第一个ticket：{issues}")
        else:
            print("✅ 测试成功！无符合条件的ticket")
    except Exception as e:
        print(f"❌ 测试失败：{str(e)}")

    issueKey = issues[1]["id"]
    print(f"{issueKey}")

    comments = jira_cli.get_issue_comments(issueKey)
    print(comments)

    attachmentKey = issues[1]["fields"]["attachment"][0]["id"]
    print(f"{attachmentKey}")
    jira_response = jira_cli.get_issue_attachments(attachmentKey)
    print(f"{jira_response}, {type(jira_response)}")
    print(f"Content-Type: {jira_response.headers.get('Content-Type')}")  # 文件类型，如 image/png, application/pdf
    print(f"Content-Length: {jira_response.headers.get('Content-Length')}")  # 文件大小（字节）
    try:
        jira_attachment = convert_response_to_json(jira_response)
    except Exception as e:
        print(e)
    print(jira_attachment)
